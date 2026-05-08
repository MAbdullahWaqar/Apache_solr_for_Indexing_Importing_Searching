"""Generate the Word (.docx) version of the Lab 13 report from the markdown.

This is intentionally a small, dependency-light Markdown -> docx renderer
tailored to the structure of `Lab_13_Report.md` (headings, paragraphs, lists,
tables, fenced code blocks). For more elaborate Markdown we would reach for
`pypandoc`; this lightweight implementation keeps the report self-contained.

Run:
    pip install -r ../backend/requirements.txt   # for python-docx
    python docs/generate_report.py               # writes Lab_13_Report.docx
"""

from __future__ import annotations

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Cm

ROOT = Path(__file__).resolve().parents[1]
REPORT_MD = ROOT / "docs" / "Lab_13_Report.md"
SCREENSHOTS_DIR = ROOT / "docs" / "screenshots"
OUTPUT_DOCX = ROOT / "docs" / "Lab_13_Report.docx"

CODE_FONT = "Consolas"
BODY_FONT = "Calibri"
HEADING_COLOR = RGBColor(0x10, 0x3A, 0x71)
CODE_BG = "F4F6FA"
TABLE_HEADER_BG = "DCE6F3"


# ---------------------------------------------------------------------------
# Low-level docx helpers
# ---------------------------------------------------------------------------
def _shade(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _add_run(paragraph, text: str, *, bold: bool = False, italic: bool = False,
             code: bool = False, color: RGBColor | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    if code:
        run.font.name = CODE_FONT
        run.font.size = Pt(10)
    else:
        run.font.name = BODY_FONT
        run.font.size = Pt(11)
    if color is not None:
        run.font.color.rgb = color


def _emit_inline(paragraph, text: str) -> None:
    """Render Markdown inline syntax: **bold**, *italic*, `code`, [text](url)."""

    pattern = re.compile(
        r"(\*\*[^*]+\*\*"
        r"|\*[^*]+\*"
        r"|`[^`]+`"
        r"|\[[^\]]+\]\([^)]+\))"
    )
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            _add_run(paragraph, text[pos:match.start()])
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            _add_run(paragraph, token[2:-2], bold=True)
        elif token.startswith("*") and token.endswith("*"):
            _add_run(paragraph, token[1:-1], italic=True)
        elif token.startswith("`") and token.endswith("`"):
            _add_run(paragraph, token[1:-1], code=True)
        elif token.startswith("["):
            link_match = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token)
            if link_match:
                _add_run(paragraph, link_match.group(1), italic=True,
                         color=RGBColor(0x10, 0x3A, 0x71))
            else:
                _add_run(paragraph, token)
        pos = match.end()
    if pos < len(text):
        _add_run(paragraph, text[pos:])


def _add_heading(doc, text: str, level: int) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    run.font.color.rgb = HEADING_COLOR
    run.font.name = BODY_FONT
    sizes = {1: 22, 2: 16, 3: 13, 4: 12, 5: 11, 6: 11}
    run.font.size = Pt(sizes.get(level, 11))
    paragraph.paragraph_format.space_before = Pt(14 if level <= 2 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    if level == 1:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_paragraph(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    _emit_inline(p, text)


def _add_bullet(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    _emit_inline(p, text)


def _add_numbered(doc, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    _emit_inline(p, text)


def _add_code_block(doc, code: str, language: str | None = None) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    _shade(cell, CODE_BG)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for line in code.splitlines() or [""]:
        run = para.add_run(line + "\n")
        run.font.name = CODE_FONT
        run.font.size = Pt(9.5)
    table.autofit = True


def _add_table(doc, header: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Light Grid Accent 1"
    hdr = table.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        _shade(hdr[i], TABLE_HEADER_BG)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _emit_inline(p, val)
            for run in p.runs:
                run.font.size = Pt(10)


def _add_horizontal_rule(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAB7C9")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Markdown parser (just enough to render Lab_13_Report.md)
# ---------------------------------------------------------------------------
def _parse_table(lines: list[str], idx: int) -> tuple[list[str], list[list[str]], int]:
    header_row = [c.strip() for c in lines[idx].strip().strip("|").split("|")]
    sep_idx = idx + 1
    body_rows: list[list[str]] = []
    j = sep_idx + 1
    while j < len(lines) and lines[j].lstrip().startswith("|"):
        cells = [c.strip() for c in lines[j].strip().strip("|").split("|")]
        if len(cells) < len(header_row):
            cells.extend([""] * (len(header_row) - len(cells)))
        body_rows.append(cells[: len(header_row)])
        j += 1
    return header_row, body_rows, j


def render_markdown(doc, md: str) -> None:
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Fenced code block
        if stripped.startswith("```"):
            language = stripped[3:].strip() or None
            j = i + 1
            buf: list[str] = []
            while j < len(lines) and not lines[j].strip().startswith("```"):
                buf.append(lines[j])
                j += 1
            _add_code_block(doc, "\n".join(buf), language)
            i = j + 1
            continue

        # Horizontal rule
        if stripped in {"---", "***", "___"}:
            _add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            text = stripped[level:].strip()
            _add_heading(doc, text, min(level, 6))
            i += 1
            continue

        # Tables
        if stripped.startswith("|") and i + 1 < len(lines) and \
           re.match(r"^\s*\|?\s*[-:]+\s*\|", lines[i + 1] or ""):
            header, body, next_i = _parse_table(lines, i)
            _add_table(doc, header, body)
            i = next_i
            continue

        # Bullet list
        if re.match(r"^\s*[-*+]\s+", line):
            text = re.sub(r"^\s*[-*+]\s+", "", line)
            _add_bullet(doc, text)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\s*\d+\.\s+", line):
            text = re.sub(r"^\s*\d+\.\s+", "", line)
            _add_numbered(doc, text)
            i += 1
            continue

        # Blank
        if not stripped:
            i += 1
            continue

        # Block quote
        if stripped.startswith(">"):
            text = stripped.lstrip(">").lstrip()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x40, 0x4A, 0x60)
            i += 1
            continue

        # Paragraph (collect continuation lines)
        buf = [stripped]
        j = i + 1
        while j < len(lines) and lines[j].strip() and \
              not re.match(r"^\s*([-*+]|\d+\.)\s+", lines[j]) and \
              not lines[j].strip().startswith(("#", "|", "```", ">")):
            buf.append(lines[j].strip())
            j += 1
        _add_paragraph(doc, " ".join(buf))
        i = j


# ---------------------------------------------------------------------------
# Cover page + screenshots appendix
# ---------------------------------------------------------------------------
def add_cover(doc) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Lab 13 — Open-Ended Lab\n"
                         "Indexing, Importing and Searching\n"
                         "data in Apache Solr")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = HEADING_COLOR

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run(
        "\nParallel and Distributed Computing\n"
        "Apache Solr 9.x · Flask · Modern Web UI\n"
    )
    s_run.font.size = Pt(13)
    s_run.font.color.rgb = RGBColor(0x4A, 0x55, 0x72)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    i_run = info.add_run(
        "\n\nMuhammad Abdullah Waqar\n"
        "https://github.com/MAbdullahWaqar/Apache_solr_for_Indexing_Importing_Searching\n"
    )
    i_run.font.size = Pt(11)

    doc.add_page_break()


def add_screenshots_appendix(doc) -> None:
    if not SCREENSHOTS_DIR.exists():
        return
    images = sorted(p for p in SCREENSHOTS_DIR.iterdir()
                    if p.suffix.lower() in {".png", ".jpg", ".jpeg"})
    if not images:
        return
    doc.add_page_break()
    _add_heading(doc, "Appendix A — Screenshots", level=1)
    for img in images:
        _add_heading(doc, img.stem, level=3)
        try:
            doc.add_picture(str(img), width=Inches(6.0))
        except Exception as exc:
            _add_paragraph(doc, f"(Could not embed {img.name}: {exc})")


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    if not REPORT_MD.exists():
        raise SystemExit(f"Report not found: {REPORT_MD}")

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)

    sections = doc.sections
    for section in sections:
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)

    add_cover(doc)
    md = REPORT_MD.read_text(encoding="utf-8")
    render_markdown(doc, md)
    add_screenshots_appendix(doc)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Wrote {OUTPUT_DOCX} ({os.path.getsize(OUTPUT_DOCX) // 1024} KB)")


if __name__ == "__main__":
    main()
